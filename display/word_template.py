import os
import numpy as np
from pathlib import Path
from docx import Document
from pandas import DataFrame
from typing import Dict, Any, List

from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class BacktestDocument:

    def __init__(self) -> None:
        self.document: Document = Document()
        self.backtest_result_dict = {}

    def get_file_name(self, params: Dict[str, Any]) -> str:
        """
        根据参数生成文件名，即将字典展开成 key1_value1 key2_value2 key3_value3 ...

        Parameters
        ----------
        params: Dict[str, Any]
                参数字典

        Returns
        -------
        string: str
                文件名 key1_value1 key2_value2 key3_value3 ...
        """
        string = ''
        for param in params:
            string += f"{param}_{params[param]}"
            string += ' '
        if string:
            string = string[:-1]
        return string

    def add_backtest_result(self, info_dict: Dict[str, Any],
                            curve = None,
                            all_metrics = None,
                            industry_metrics = None,
                            symbol_metrics = None,
                            cum_profit_series = None) -> None:
        """
        添加回测结果信息，包括参数字典, 曲线路径, all_metrics, industry_metrics, symbol_metrics, cum_profit_series

        Parameters
        ----------
        info_dict: Dict[str, Any]
                    参数字典

        curve: str
                曲线图片文件路径

        all_metrics: DataFramne
                    回测整体的指标

        industry_metrics: DataFrame
                          按行业分的指标

        symbol_metrics: DataFrame
                        按品种分的指标
        cum_profit_series: DataFrame
                            累计收益序列，第一列为时间，第二列为累计收益

        Returns
        -------
        None
        """

        str_info_dict = str(info_dict)
        self.backtest_result_dict[str_info_dict] = {'curve':curve,
                                                    'all': all_metrics,
                                                    'industry': industry_metrics,
                                                    'symbol': symbol_metrics,
                                                    'cum_profit_series': cum_profit_series}

    def display_backtest_result(self, info_list: List[Dict], output_folder_path: str, doc_name: str) -> None:
        """
        输出回测结果，生成word文档

        Parameters
        ----------
        info_list: str,
                    多组参数列表

        output_folder_path: str
                            输出文件夹路径

        doc_name: str
                  输出文件名称

        Returns
        -------
        None
        """
        output_folder_path = Path(output_folder_path)
        document = self.document
        document.styles['Normal'].font.name = u'Times New Roman'
        document.add_heading('backtest result', level=0)
        for info in info_list:

            i = info_list.index(info)+1
            single_output_folder_path = output_folder_path.joinpath(str(i))
            if not os.path.exists(single_output_folder_path):
                os.makedirs(single_output_folder_path)
            document.add_heading(str(i), level=1)
            p1 = document.add_paragraph()
            # 添加参数说明
            param_string = ''
            for key, value in info.items():
                param_string += f"{key}={str(value)}\n"
            run1 = p1.add_run(param_string)

            # 添加回测结果
            backtest_result = self.backtest_result_dict[str(info)]

            # 添加资金曲线
            p2 = document.add_paragraph()
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            curve = backtest_result['curve']
            if not isinstance(curve, str):
                raise ValueError("curve no found!")
            run2 = p2.add_run("")
            run2.add_picture(str(curve),width = Inches(7))

            # 添加metrics表格
            p3 = document.add_paragraph()
            p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            all_metrics = np.round(backtest_result['all'], 2)
            if isinstance(all_metrics, DataFrame):
                all_metrics.to_csv(single_output_folder_path.joinpath("symbol.csv"))
                table3 = document.add_table(rows=all_metrics.shape[0], cols=all_metrics.shape[1])
                table3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for i in range(all_metrics.shape[0]):
                    for j in range(all_metrics.shape[1]):
                        table3.cell(i,j).text = str(all_metrics.iloc[i,j])

            industry_metrics = np.round(backtest_result['industry'], 2)
            if isinstance(industry_metrics, DataFrame):
                industry_metrics.to_csv(single_output_folder_path.joinpath("industry.csv"))

            symbol_metrics = np.round(backtest_result['symbol'], 2)
            if isinstance(symbol_metrics, DataFrame):
                symbol_metrics.to_csv(single_output_folder_path.joinpath("symbol.csv"))

        # 保存word文档
        document.save(output_folder_path.joinpath(doc_name))




